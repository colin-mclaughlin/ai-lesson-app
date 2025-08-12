from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from typing import Union, List, Optional
import logging
import re
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO

# Import our existing lesson generation functions
from prompt_builder import build_grammar_lesson_prompt, build_multi_rule_grammar_lesson_prompt
from openai_client import generate_lesson
from database import LessonDatabase

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def clean_lesson_text(text: str, subject: str, topic: str) -> str:
    """
    Clean lesson text by removing markdown formatting, placeholder headings,
    and adding a clean title.
    """
    if not text:
        return text
    
    # Remove markdown formatting first
    # Replace **bold** → bold
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    # Replace _italics_ → italics
    text = re.sub(r'_(.*?)_', r'\1', text)
    # Remove heading syntax (##, ###)
    text = re.sub(r'^#{1,3}\s*', '', text, flags=re.MULTILINE)
    
    cleaned_lines = []
    activity_counter = 0
    rule_section_processed = False
    title_added = False
    
    # Regex patterns (compiled for efficiency and case-insensitivity)
    rule_heading_pattern = re.compile(r'^\s*Rule Heading\s*$', re.IGNORECASE)
    rule_x_pattern = re.compile(r'^\s*Rule \d+:\s*.*$', re.IGNORECASE)
    # This pattern now captures optional leading numbers for activity sections
    activity_section_pattern = re.compile(r'^\s*(?:\d+\.\s*)?(Activity|Section) Section [A-Z]\s*$', re.IGNORECASE)
    bullet_pattern = re.compile(r'^\s*[-•]\s*') # For removing leading bullets
    topic_pattern = re.compile(r'^\s*' + re.escape(topic) + r'\s*$', re.IGNORECASE)
    
    lines = text.split('\n')
    
    # Add clean title at the top
    cleaned_lines.append(f"{subject} — {topic}")
    cleaned_lines.append("") # Add a blank line after title
    title_added = True
    
    for line in lines:
        # 1. Handle "Rule Heading" and "Rule X:" replacement
        if rule_heading_pattern.match(line) or rule_x_pattern.match(line):
            if not rule_section_processed:
                # Don't add topic again if it's already in the title
                if not title_added:
                    cleaned_lines.append(topic)
                    cleaned_lines.append("") # Blank line after topic
                cleaned_lines.append("Explanation")
                cleaned_lines.append("") # Blank line after explanation
                rule_section_processed = True
            continue # Skip the original rule line
        
        # 2. Remove any standalone line that equals the topic (case-insensitive) if it appears again near the top
        if topic_pattern.match(line) and title_added:
            continue # Skip duplicate topic lines
        
        # 3. Handle Activity Section renaming and remove leading numbers
        match_activity = activity_section_pattern.match(line)
        if match_activity:
            activity_counter += 1
            cleaned_lines.append(f"Activity {activity_counter}")
            cleaned_lines.append("") # Blank line after activity title
            continue
        
        # 4. Remove leading "-" or "•" from list items
        line = bullet_pattern.sub('', line)
        
        # Add the processed line
        cleaned_lines.append(line)
    
    # 5. Ensure the cleaned text is returned without any extra blank lines caused by these removals.
    # Trim extra blank lines (more than 2 consecutive empty lines -> 1 empty line)
    final_cleaned_lines = []
    consecutive_blanks = 0
    for line in cleaned_lines:
        if not line.strip():
            consecutive_blanks += 1
        else:
            consecutive_blanks = 0
        
        if consecutive_blanks <= 1: # Allow at most one blank line
            final_cleaned_lines.append(line)
    
    # Remove leading/trailing blank lines from the final list
    while final_cleaned_lines and not final_cleaned_lines[0].strip():
        final_cleaned_lines.pop(0)
    while final_cleaned_lines and not final_cleaned_lines[-1].strip():
        final_cleaned_lines.pop()
    
    return '\n'.join(final_cleaned_lines)


def parse_lesson_sections(text: str) -> dict:
    """
    Parse lesson text into sections: explanation and activities.
    Returns {"explanation": str, "activities": {1: str, 2: str, 3: str, 4: str}}
    """
    sections = {"explanation": "", "activities": {}}
    
    # Split text into lines
    lines = text.split('\n')
    current_section = "explanation"
    current_content = []
    
    for line in lines:
        # Check for Activity headers
        activity_match = re.match(r'^Activity\s+(\d+)\s*$', line.strip(), re.IGNORECASE)
        if activity_match:
            # Save previous section
            if current_section == "explanation":
                sections["explanation"] = '\n'.join(current_content).strip()
            else:
                sections["activities"][int(current_section)] = '\n'.join(current_content).strip()
            
            # Start new activity section
            current_section = activity_match.group(1)
            current_content = []
        else:
            current_content.append(line)
    
    # Save the last section
    if current_section == "explanation":
        sections["explanation"] = '\n'.join(current_content).strip()
    else:
        sections["activities"][int(current_section)] = '\n'.join(current_content).strip()
    
    return sections


def count_numbered_items(block: str) -> int:
    """
    Count numbered items in a text block using regex pattern.
    """
    if not block:
        return 0
    
    # Find all lines that match the pattern: start with optional whitespace, then number, dot, space
    matches = re.findall(r'^\s*\d+\.\s', block, re.MULTILINE)
    return len(matches)


def validate_counts(sections: dict, config: dict) -> dict:
    """
    Validate that each activity has the expected number of items.
    Returns dict with per-activity actual vs expected counts and overall 'ok' status.
    """
    validation_result = {
        "ok": True,
        "activities": {},
        "summary": {}
    }
    
    expected_counts = {
        1: config.get("section_a_questions", 6),
        2: config.get("section_b_questions", 6),
        3: config.get("section_c_questions", 6),
        4: config.get("section_d_questions", 6)
    }
    
    for activity_num in range(1, 5):
        activity_content = sections["activities"].get(activity_num, "")
        actual_count = count_numbered_items(activity_content)
        expected_count = expected_counts[activity_num]
        
        validation_result["activities"][activity_num] = {
            "actual": actual_count,
            "expected": expected_count,
            "match": actual_count == expected_count
        }
        
        if actual_count != expected_count:
            validation_result["ok"] = False
    
    # Add summary
    validation_result["summary"] = {
        "total_activities": len(sections["activities"]),
        "matching_activities": sum(1 for act in validation_result["activities"].values() if act["match"]),
        "total_expected": sum(expected_counts.values()),
        "total_actual": sum(act["actual"] for act in validation_result["activities"].values())
    }
    
    return validation_result


def generate_targeted_activity_prompt(topic: str, activity_num: int, expected_count: int) -> str:
    """
    Generate a targeted prompt for regenerating a specific activity.
    """
    return f"""Regenerate ONLY Activity {activity_num} for topic "{topic}", exactly {expected_count} items, each on its own line starting '1.' to '{expected_count}.' No extra text.

Example format:
1. [First item]
2. [Second item]
3. [Third item]
...
{expected_count}. [Last item]

Topic: {topic}
Activity {activity_num}:"""


def fix_activity_section(lesson_text: str, topic: str, activity_num: int, expected_count: int, max_attempts: int = 2) -> str:
    """
    Attempt to fix a specific activity section by regenerating it.
    Returns the updated lesson text.
    """
    logger.info(f"Attempting to fix Activity {activity_num} (expected: {expected_count} items)")
    
    for attempt in range(max_attempts):
        try:
            # Generate targeted prompt
            targeted_prompt = generate_targeted_activity_prompt(topic, activity_num, expected_count)
            
            # Generate new activity content
            new_activity_content = generate_lesson(targeted_prompt)
            
            # Clean up the generated content (remove any extra text)
            lines = new_activity_content.split('\n')
            cleaned_lines = []
            for line in lines:
                if re.match(r'^\s*\d+\.\s', line.strip()):
                    cleaned_lines.append(line)
            
            if len(cleaned_lines) == expected_count:
                # Replace the activity section in the original text
                sections = parse_lesson_sections(lesson_text)
                sections["activities"][activity_num] = '\n'.join(cleaned_lines)
                
                # Reconstruct the lesson text
                updated_text = sections["explanation"] + "\n\n"
                for act_num in range(1, 5):
                    if act_num in sections["activities"]:
                        updated_text += f"Activity {act_num}\n{sections['activities'][act_num]}\n\n"
                
                logger.info(f"Successfully fixed Activity {activity_num} on attempt {attempt + 1}")
                return updated_text.strip()
            else:
                logger.warning(f"Activity {activity_num} attempt {attempt + 1}: got {len(cleaned_lines)} items, expected {expected_count}")
                
        except Exception as e:
            logger.error(f"Error fixing Activity {activity_num} attempt {attempt + 1}: {str(e)}")
    
    # If all attempts failed, try to synthesize the missing items
    logger.warning(f"All attempts failed for Activity {activity_num}, synthesizing items")
    return synthesize_missing_items(lesson_text, topic, activity_num, expected_count)


def synthesize_missing_items(lesson_text: str, topic: str, activity_num: int, expected_count: int) -> str:
    """
    Synthesize missing items for an activity when regeneration fails.
    """
    sections = parse_lesson_sections(lesson_text)
    current_content = sections["activities"].get(activity_num, "")
    current_count = count_numbered_items(current_content)
    
    if current_count >= expected_count:
        return lesson_text
    
    # Extract existing items
    lines = current_content.split('\n')
    existing_items = []
    for line in lines:
        if re.match(r'^\s*\d+\.\s', line.strip()):
            existing_items.append(line.strip())
    
    # Generate additional items
    missing_count = expected_count - current_count
    additional_prompt = f"""Generate exactly {missing_count} more numbered items for Activity {activity_num} about "{topic}". 
Each item should be on its own line starting with the next number ({current_count + 1}.) through ({expected_count}.).
Make them relevant to the topic and consistent with the existing items.

Existing items:
{chr(10).join(existing_items[:3])}  # Show first 3 for context

Generate {missing_count} more items:"""
    
    try:
        additional_content = generate_lesson(additional_prompt)
        additional_lines = additional_content.split('\n')
        additional_items = []
        
        for line in additional_lines:
            if re.match(r'^\s*\d+\.\s', line.strip()):
                additional_items.append(line.strip())
                if len(additional_items) >= missing_count:
                    break
        
        # Combine existing and additional items
        all_items = existing_items + additional_items[:missing_count]
        
        # Update the activity section
        sections["activities"][activity_num] = '\n'.join(all_items)
        
        # Reconstruct the lesson text
        updated_text = sections["explanation"] + "\n\n"
        for act_num in range(1, 5):
            if act_num in sections["activities"]:
                updated_text += f"Activity {act_num}\n{sections['activities'][act_num]}\n\n"
        
        logger.info(f"Synthesized {len(additional_items[:missing_count])} items for Activity {activity_num}")
        return updated_text.strip()
        
    except Exception as e:
        logger.error(f"Error synthesizing items for Activity {activity_num}: {str(e)}")
        return lesson_text


app = FastAPI(title="Coding Cat Lesson Generator API", version="1.0.0")

# Configure CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000", "http://localhost:5173", "http://127.0.0.1:3000", "http://127.0.0.1:5173"],  # React dev server ports
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Request model
class LessonRequest(BaseModel):
    grade: Union[int, str]
    subject: str
    topic: str
    questions_per_section: int = 6

# Response model
class LessonResponse(BaseModel):
    lessonText: str
    success: bool = True
    regenerated: bool = False
    warnings: List[str] = []
    lessonId: Optional[int] = None

# Lesson history models
class LessonSummary(BaseModel):
    id: int
    topics: List[str]
    grade: int
    age: Optional[int]
    date_generated: str

class LessonsListResponse(BaseModel):
    lessons: List[LessonSummary]
    total: int

# Initialize database
db = LessonDatabase()

def validate_lesson(text: str, topic: str) -> tuple[bool, List[str]]:
    """
    Validate lesson content for banned terms and topic adherence.
    Returns (is_valid, warnings_list)
    """
    warnings = []
    
    # Check for banned terms (pictures, drawings, diagrams, images)
    banned_pattern = r"\b(picture|draw|diagram|image|illustration)\b"
    banned_matches = re.findall(banned_pattern, text, flags=re.IGNORECASE)
    if banned_matches:
        warnings.append(f"Found banned terms: {', '.join(set(banned_matches))}")
    
    # Check topic adherence (topic should appear at least 2 times)
    topic_matches = len(re.findall(re.escape(topic), text, flags=re.IGNORECASE))
    if topic_matches < 2:
        warnings.append(f"Topic '{topic}' only appears {topic_matches} times")
    
    is_valid = len(warnings) == 0
    return is_valid, warnings

def create_docx_from_lesson(lesson_text: str, grade: int, topics: List[str]) -> BytesIO:
    """
    Create a Word document from lesson text with proper styling.
    """
    doc = Document()
    
    # Set up document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Add header
    header = doc.add_heading('Coding Cat Club', 0)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add grade and topic info
    topic_text = ', '.join(topics)
    grade_topic = doc.add_paragraph(f'Grade {grade} — Topic: {topic_text}')
    grade_topic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add Name and Date lines
    name_date = doc.add_paragraph()
    name_date.add_run('Name: ').bold = True
    name_date.add_run('_' * 30)
    name_date.add_run('    Date: ').bold = True
    name_date.add_run('_' * 30)
    
    # Add title
    title = doc.add_heading('Lesson Worksheet', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Process lesson text line by line
    lines = lesson_text.split('\n')
    i = 0
    
    # Skip the first line if it's the subject-topic title (already handled in header)
    if lines and ' — ' in lines[0]:
        i = 2  # Skip title and empty line
    
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            # Empty line - add paragraph break
            doc.add_paragraph()
            i += 1
            continue
        
        # Check for section headers (previously bold text, now plain text)
        # Look for lines that are likely section headers (all caps, short, or end with colon)
        if (line.isupper() and len(line) < 50) or line.endswith(':') and len(line) < 100:
            # Section header
            doc.add_heading(line, level=2)
            i += 1
            continue
        
        # Check for numbered lists
        numbered_match = re.match(r'^(\d+)\.\s*(.+)$', line)
        if numbered_match:
            # Start numbered list
            list_items = []
            while i < len(lines) and re.match(r'^\d+\.\s*(.+)$', lines[i].strip()):
                match = re.match(r'^\d+\.\s*(.+)$', lines[i].strip())
                if match:
                    list_items.append(match.group(1))
                i += 1
            
            # Add numbered list with manual numbering (restarts at 1 for each block)
            for idx, item in enumerate(list_items, 1):
                p = doc.add_paragraph(f"{idx}. {item}")
                # Check if next line is an answer line
                if i < len(lines) and lines[i].strip() == '':
                    i += 1
                    if i < len(lines) and lines[i].strip() == '':
                        # Add answer line
                        answer_p = doc.add_paragraph()
                        answer_run = answer_p.add_run('_' * 50)
                        answer_run.underline = True
                        i += 1
            continue
        
        # Check for bullet points
        bullet_match = re.match(r'^[-•]\s*(.+)$', line)
        if bullet_match:
            # Start bullet list
            list_items = []
            while i < len(lines) and re.match(r'^[-•]\s*(.+)$', lines[i].strip()):
                match = re.match(r'^[-•]\s*(.+)$', lines[i].strip())
                if match:
                    list_items.append(match.group(1))
                i += 1
            
            # Add bullet list
            for item in list_items:
                doc.add_paragraph(item, style='List Bullet')
            continue
        
        # Check for instructions
        if line.startswith('Instructions:'):
            p = doc.add_paragraph()
            p.add_run('Instructions: ').bold = True
            p.add_run(line.replace('Instructions:', '').strip()).italic = True
            i += 1
            continue
        
        # Regular paragraph with markdown processing
        p = doc.add_paragraph()
        
        # Process bold text within the line
        parts = re.split(r'(\*\*.*?\*\*)', line)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                # Bold text
                run = p.add_run(part.replace('**', ''))
                run.bold = True
            else:
                # Regular text
                p.add_run(part)
        
        # Check if next line should be an answer line
        if i + 1 < len(lines) and lines[i + 1].strip() == '':
            i += 1
            if i + 1 < len(lines) and lines[i + 1].strip() == '':
                # Add answer line
                answer_p = doc.add_paragraph()
                answer_run = answer_p.add_run('_' * 50)
                answer_run.underline = True
                i += 1
        
        i += 1
    
    # Save to BytesIO
    docx_buffer = BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    return docx_buffer

@app.get("/")
async def root():
    return {"message": "Coding Cat Lesson Generator API", "status": "running"}

@app.get("/health")
async def health_check():
    return {"status": "healthy", "message": "API is running"}

@app.get("/api/lessons/{lesson_id}/docx")
async def download_lesson_docx(lesson_id: int):
    """
    Download a lesson as a Word document (.docx)
    """
    try:
        logger.info(f"Generating DOCX for lesson ID: {lesson_id}")
        
        # Get lesson from database
        lesson = db.get_lesson(lesson_id)
        if not lesson:
            raise HTTPException(status_code=404, detail=f"Lesson with ID {lesson_id} not found")
        
        # Create DOCX
        docx_buffer = create_docx_from_lesson(
            lesson['lesson_text'],
            lesson['grade'],
            lesson['topics']
        )
        
        # Create filename
        topic_text = '_'.join(lesson['topics']).replace(' ', '_')
        filename = f"Lesson_Grade{lesson['grade']}_{topic_text}.docx"
        
        logger.info(f"DOCX generated successfully: {filename}")
        
        return StreamingResponse(
            docx_buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error generating DOCX for lesson {lesson_id}: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to generate DOCX: {str(e)}")

@app.post("/api/generate-lesson", response_model=LessonResponse)
async def generate_lesson_endpoint(request: LessonRequest):
    """
    Generate a lesson based on the provided parameters.
    """
    try:
        logger.info(f"Generating lesson for grade {request.grade}, subject {request.subject}, topic {request.topic}")
        
        # Convert grade to int if it's a string
        grade_level = int(request.grade) if isinstance(request.grade, str) else request.grade
        
        # Validate grade level (expanded to 1-12)
        if grade_level < 1 or grade_level > 12:
            raise HTTPException(status_code=422, detail="Grade level must be between 1 and 12")
        
        # Validate questions per section
        if request.questions_per_section < 1 or request.questions_per_section > 20:
            raise HTTPException(status_code=400, detail="Questions per section must be between 1 and 20")
        
        # Create lesson configuration
        lesson_config = {
            "grade_level": grade_level,
            "section_a_questions": request.questions_per_section,
            "section_b_questions": request.questions_per_section,
            "section_c_questions": request.questions_per_section,
            "section_d_questions": request.questions_per_section
        }
        
        # For now, we'll focus on single topic lessons
        # In the future, we could split the topic by commas or other delimiters for multi-topic lessons
        topics = [topic.strip() for topic in request.topic.split(',') if topic.strip()]
        
        if not topics:
            raise HTTPException(status_code=400, detail="At least one topic must be provided")
        
        # Generate the prompt based on number of topics
        if len(topics) == 1:
            prompt = build_grammar_lesson_prompt(topics[0], lesson_config)
        else:
            prompt = build_multi_rule_grammar_lesson_prompt(topics, lesson_config)
        
        logger.info("Calling OpenAI API to generate lesson...")
        
        # Generate the lesson using OpenAI
        lesson_text = generate_lesson(prompt)
        
        # Validate the generated lesson
        is_valid, warnings = validate_lesson(lesson_text, topics[0])
        
        # If invalid, regenerate once with stronger constraints
        regenerated = False
        if not is_valid:
            logger.warning(f"Lesson validation failed: {warnings}. Regenerating with stronger constraints.")
            
            # Add stronger system steer to the prompt
            stronger_prompt = prompt + f"\n\nYour last output violated constraints. Strictly follow: no pictures; keep strictly on-topic: {topics[0]}."
            
            # Regenerate
            lesson_text = generate_lesson(stronger_prompt)
            regenerated = True
            
            # Validate again
            is_valid, new_warnings = validate_lesson(lesson_text, topics[0])
            warnings = new_warnings  # Use the new warnings from regeneration
        
        logger.info("Lesson generated successfully")
        
        # Clean the lesson text
        cleaned_lesson_text = clean_lesson_text(lesson_text, request.subject, topics[0])
        
        # Validate activity counts and fix if needed
        sections = parse_lesson_sections(cleaned_lesson_text)
        validation_result = validate_counts(sections, lesson_config)
        
        # Log initial validation results
        logger.info(f"Initial validation: {validation_result['summary']}")
        for act_num, act_data in validation_result['activities'].items():
            if not act_data['match']:
                logger.warning(f"Activity {act_num}: expected {act_data['expected']}, got {act_data['actual']}")
        
        # Fix mismatched activities
        if not validation_result['ok']:
            logger.info("Fixing mismatched activities...")
            fixed_lesson_text = cleaned_lesson_text
            
            for act_num, act_data in validation_result['activities'].items():
                if not act_data['match']:
                    logger.info(f"Fixing Activity {act_num}: {act_data['actual']} -> {act_data['expected']} items")
                    fixed_lesson_text = fix_activity_section(
                        fixed_lesson_text, 
                        topics[0], 
                        act_num, 
                        act_data['expected']
                    )
            
            # Re-validate after fixes
            fixed_sections = parse_lesson_sections(fixed_lesson_text)
            final_validation = validate_counts(fixed_sections, lesson_config)
            
            # Log final validation results
            logger.info(f"Final validation: {final_validation['summary']}")
            for act_num, act_data in final_validation['activities'].items():
                if not act_data['match']:
                    logger.warning(f"Activity {act_num} still mismatched: expected {act_data['expected']}, got {act_data['actual']}")
                else:
                    logger.info(f"Activity {act_num} fixed: {act_data['actual']} items")
            
            cleaned_lesson_text = fixed_lesson_text
        
        # Save lesson to database (optional - you can remove this if you don't want to auto-save)
        try:
            lesson_id = db.save_lesson(topics, grade_level, cleaned_lesson_text)
            logger.info(f"Lesson saved to database with ID: {lesson_id}")
            print(f"DEBUG: Lesson saved with ID: {lesson_id}")  # Debug log
        except Exception as e:
            logger.warning(f"Failed to save lesson to database: {e}")
            lesson_id = None  # Set to None if save fails
            # Don't fail the request if database save fails
        
        return LessonResponse(
            lessonText=cleaned_lesson_text,
            regenerated=regenerated,
            warnings=warnings,
            lessonId=lesson_id
        )
        
    except HTTPException:
        # Re-raise HTTP exceptions
        raise
    except Exception as e:
        logger.error(f"Error generating lesson: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to generate lesson: {str(e)}")

@app.get("/api/lessons", response_model=LessonsListResponse)
async def get_lessons_list():
    """
    Get all lessons from the database, ordered by date descending.
    Returns a list of lesson summaries with basic information.
    """
    try:
        logger.info("Fetching all lessons from database")
        lessons = db.list_lessons()
        
        # Convert to response format
        lesson_summaries = [
            LessonSummary(
                id=lesson['id'],
                topics=lesson['topics'],
                grade=lesson['grade'],
                age=lesson['age'],
                date_generated=lesson['date_generated']
            )
            for lesson in lessons
        ]
        
        return LessonsListResponse(
            lessons=lesson_summaries,
            total=len(lesson_summaries)
        )
        
    except Exception as e:
        logger.error(f"Error fetching lessons: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to fetch lessons: {str(e)}")

@app.get("/api/lessons/{lesson_id}")
async def get_lesson_by_id(lesson_id: int):
    """
    Get a specific lesson by ID, including full lesson content.
    Returns 404 if lesson not found.
    """
    try:
        logger.info(f"Fetching lesson with ID: {lesson_id}")
        lesson = db.get_lesson(lesson_id)
        
        if not lesson:
            raise HTTPException(status_code=404, detail=f"Lesson with ID {lesson_id} not found")
        
        return lesson
        
    except HTTPException:
        # Re-raise HTTP exceptions
        raise
    except Exception as e:
        logger.error(f"Error fetching lesson {lesson_id}: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to fetch lesson: {str(e)}")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)